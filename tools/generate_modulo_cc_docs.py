"""
Generate .docx files for Módulo Claude Code course support docs.
Uses python-docx with native H1/H2/H3/H4 headings — no ASCII dividers.
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path(r"c:\runa-systems-global\SÍRIOS\📦 Entregáveis\runa-systems-skool\modulo-claude-code")
OUTPUT = BASE

DOCS = [
    "claude-code-01-instalacao.md",
    "claude-code-02-principios-basicos.md",
    "claude-code-03-skills-plugins-mcps-clis.md",
]


def add_table_from_md(doc, lines, start_idx):
    """Parse markdown table lines and add as Word table."""
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i])
        i += 1

    # Filter separator rows
    rows = [l for l in table_lines if not re.match(r"\s*\|[\s\-:|]+\|\s*$", l)]
    if not rows:
        return i

    # Parse cells
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    if not parsed:
        return i

    ncols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=ncols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(parsed):
        for c_idx, cell_text in enumerate(row):
            if c_idx < ncols:
                cell = table.cell(r_idx, c_idx)
                # Strip markdown bold/italic
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
                clean = re.sub(r"`(.+?)`", r"\1", clean)
                clean = re.sub(r"\*(.+?)\*", r"\1", clean)
                clean = re.sub(r"\[\[(.+?)\]\]", r"\1", clean)
                p = cell.paragraphs[0]
                run = p.add_run(clean)
                if r_idx == 0:
                    run.bold = True

    doc.add_paragraph()
    return i


def add_code_block(doc, code_lines):
    """Add a code block as styled paragraph."""
    code_text = "\n".join(code_lines)
    p = doc.add_paragraph()
    p.style = "No Spacing"
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(18)


def clean_inline(text):
    """Remove markdown inline formatting."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"\[\[(.+?)\]\]", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


def md_to_docx(md_path: Path, docx_path: Path):
    content = md_path.read_text(encoding="utf-8")
    lines = content.split("\n")

    doc = Document()

    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    i = 0
    # Skip YAML frontmatter
    if lines and lines[0].strip() == "---":
        i = 1
        while i < len(lines) and lines[i].strip() != "---":
            i += 1
        i += 1  # skip closing ---

    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        raw = line.rstrip()

        # Code block start/end
        if raw.startswith("```"):
            if in_code:
                if code_lines:
                    add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.+)$", raw)
        if m:
            level = len(m.group(1))
            text = clean_inline(m.group(2))
            heading_style = f"Heading {level}"
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Horizontal rule — skip
        if re.match(r"^---+$", raw):
            i += 1
            continue

        # Table
        if raw.strip().startswith("|"):
            i = add_table_from_md(doc, lines, i)
            continue

        # Blockquote
        if raw.startswith(">"):
            text = clean_inline(raw.lstrip("> ").strip())
            if text:
                p = doc.add_paragraph(text, style="Quote")
            i += 1
            continue

        # Bullet list
        bm = re.match(r"^(\s*)[*\-]\s+(.+)$", raw)
        if bm:
            indent = len(bm.group(1)) // 2
            text = clean_inline(bm.group(2))
            p = doc.add_paragraph(text, style="List Bullet")
            if indent > 0:
                p.paragraph_format.left_indent = Pt(18 * (indent + 1))
            i += 1
            continue

        # Numbered list
        nm = re.match(r"^\s*\d+\.\s+(.+)$", raw)
        if nm:
            text = clean_inline(nm.group(1))
            doc.add_paragraph(text, style="List Number")
            i += 1
            continue

        # Empty line
        if not raw.strip():
            i += 1
            continue

        # Normal paragraph
        text = clean_inline(raw)
        if text:
            doc.add_paragraph(text)
        i += 1

    doc.save(str(docx_path))
    print(f"  OK: {docx_path.name}")


if __name__ == "__main__":
    print("Generating .docx files for Módulo Claude Code...")
    for md_name in DOCS:
        md_path = BASE / md_name
        docx_name = md_name.replace(".md", ".docx")
        docx_path = OUTPUT / docx_name
        if not md_path.exists():
            print(f"  SKIP (not found): {md_name}")
            continue
        try:
            md_to_docx(md_path, docx_path)
        except Exception as e:
            print(f"  ERROR {md_name}: {e}")
            import traceback; traceback.print_exc()
    print("Done.")
