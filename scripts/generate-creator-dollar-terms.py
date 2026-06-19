"""
Generate CREATOR$ Terms of Use as .docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = "D:/Runa/runa-systems-global/SÍRIOS/📦 Entregáveis/creator-dollar-skool-docx"
os.makedirs(OUT_DIR, exist_ok=True)


def new_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    return doc


def h1(doc, text): doc.add_heading(text, level=1)
def h2(doc, text): doc.add_heading(text, level=2)
def h3(doc, text): doc.add_heading(text, level=3)
def p(doc, text): doc.add_paragraph(text)
def blank(doc): doc.add_paragraph("")
def bullet(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def generate():
    doc = new_doc()

    h1(doc, "Termos de Uso — CREATOR$")
    p(doc, "Runa Systems Global · Versão 1.0 · Abril 2026")

    h2(doc, "Aceitação dos Termos")
    p(doc, "Ao acessar e utilizar o programa CREATOR$, você concorda integralmente com os termos descritos neste documento.")
    p(doc, "Se não concordar com qualquer parte destes termos, não utilize o programa.")

    h2(doc, "Natureza do Produto")
    p(doc, "CREATOR$ é um programa de educação digital que ensina o método de criação de avatares visuais com inteligência artificial — incluindo geração de imagens, vídeos e o Storyboard permanente da personagem.")
    p(doc, "O programa é composto por:")
    bullet(doc, [
        "Aulas em vídeo com demonstrações ao vivo (build in public)",
        "Documentos de apoio por módulo",
        "Agentes de IA incluídos (LENS e REEL)",
        "Templates e Prompt Pack",
        "Acesso à comunidade na plataforma Skool",
    ])
    p(doc, "CREATOR$ é um programa educacional. Ele entrega conhecimento, método e ferramentas. O que você faz com esse conhecimento é de sua inteira responsabilidade.")

    h2(doc, "Responsabilidades")

    h3(doc, "A Runa Systems se responsabiliza por")
    bullet(doc, [
        "Fornecer acesso completo ao conteúdo do programa na plataforma Skool",
        "Disponibilizar os agentes LENS e REEL durante a vigência do acesso",
        "Suporte técnico em caso de falhas na plataforma",
        "Atualizações de conteúdo sem custo adicional quando houver melhorias no método",
    ])

    h3(doc, "Você se responsabiliza por")
    bullet(doc, [
        "Aplicar o conhecimento adquirido — o programa ensina, mas a execução é sua",
        "O conteúdo que você gera e publica com os avatares criados",
        "A veracidade e ética do que você comunica usando o avatar",
        "Conformidade legal com as ferramentas utilizadas (termos de uso do Nano Banana Pro, Higgsfield, Sora 2 e demais ferramentas ensinadas)",
        "Impostos, obrigações fiscais e aspectos jurídicos do seu negócio",
        "Atendimento aos seus clientes e qualidade do que você vende",
        "Resultados financeiros decorrentes do uso do conhecimento",
    ])

    h2(doc, "Uso Responsável do Conhecimento e das Ferramentas de IA")
    p(doc, "CREATOR$ ensina técnicas avançadas de geração de imagens e vídeos com inteligência artificial. Este conhecimento carrega responsabilidade ética. Ao utilizar o método aprendido, você concorda com as seguintes diretrizes:")

    h3(doc, "Permitido")
    bullet(doc, [
        "Criar avatares fictícios originais para uso no seu próprio negócio ou conteúdo",
        "Criar uma versão estilizada de si mesmo como persona de conteúdo",
        "Usar o avatar para produzir conteúdo educativo, comercial ou de entretenimento ético",
        "Aplicar o método para clientes seus (se você é consultor ou agência)",
        "Ensinar o método a pessoas que adquiriram acesso próprio ao programa",
    ])

    h3(doc, "Proibido")
    bullet(doc, [
        "Criar avatares que imitem, reproduzam ou se passem por pessoas reais sem consentimento explícito",
        "Gerar deepfakes ou conteúdo que induza o público a acreditar que uma pessoa real disse ou fez algo que não disse ou fez",
        "Usar o avatar para criar ou promover conteúdo fraudulento, enganoso ou que viole direitos de terceiros",
        "Criar personas fictícias para aplicar golpes, esquemas de pirâmide ou promessas financeiras falsas",
        "Usar as técnicas aprendidas para gerar conteúdo que viole leis de proteção ao consumidor, direitos autorais ou privacidade",
        "Compartilhar os materiais do programa (aulas, documentos, templates, prompts) fora da plataforma sem autorização",
        "Vender acesso compartilhado ao programa — cada aluno precisa de acesso próprio",
        "Reproduzir o método completo como produto próprio sem licença",
    ])
    p(doc, "O uso indevido do conhecimento adquirido é de responsabilidade exclusiva do aluno. A Runa Systems não se responsabiliza por conteúdos gerados, publicados ou comercializados pelos alunos após o acesso ao programa.")

    h2(doc, "Limitações e Isenções de Responsabilidade")
    p(doc, "CREATOR$ entrega um método validado. Os resultados, no entanto, dependem de fatores que estão fora do nosso controle:")
    bullet(doc, [
        "Sua consistência de execução — aprender e não aplicar não gera resultado",
        "A qualidade da sua estratégia de conteúdo e posicionamento",
        "As políticas e atualizações das ferramentas de IA de terceiros (Nano Banana Pro, Higgsfield, Sora 2, etc.)",
        "Mudanças nos algoritmos das plataformas de distribuição (Instagram, YouTube, etc.)",
        "Fatores de mercado, nicho escolhido e contexto econômico",
    ])
    p(doc, "A Runa Systems não se responsabiliza por:")
    bullet(doc, [
        "Prejuízos financeiros decorrentes de decisões tomadas com base no conteúdo do programa",
        "Falhas técnicas ou mudanças nas ferramentas de terceiros ensinadas no curso",
        "Resultados abaixo do esperado por falta de execução ou aplicação incorreta do método",
        "Conteúdo gerado pelos alunos e suas consequências jurídicas, éticas ou comerciais",
        "Problemas de acesso à plataforma Skool por questões técnicas fora do nosso controle",
    ])

    h2(doc, "Propriedade Intelectual")

    h3(doc, "O que é seu")
    bullet(doc, [
        "Todo avatar, imagem, vídeo e Storyboard que você criar aplicando o método é de sua propriedade",
        "O conteúdo gerado a partir da sua persona e identidade é seu",
        "Você pode usar comercialmente sem restrição tudo que criar com o método",
    ])

    h3(doc, "O que é da Runa Systems")
    bullet(doc, [
        "O programa CREATOR$ e sua estrutura pedagógica",
        "Os agentes LENS e REEL (GPTs incluídos no programa)",
        "Os templates, documentos de apoio e Prompt Pack",
        "A metodologia Runa de construção de avatares",
    ])
    p(doc, "Você não pode reproduzir, revender ou redistribuir os materiais do programa sem autorização expressa por escrito da Runa Systems.")

    h2(doc, "Política de Reembolso")

    h3(doc, "Garantia de 7 dias")
    p(doc, "Se você acessar o CREATOR$, assistir às aulas e decidir que o programa não é para você dentro de 7 dias da compra:")
    bullet(doc, [
        "Devolução de 100% do valor pago",
        "Sem questionamentos",
        "Envio de e-mail solicitando reembolso dentro de 7 dias da data de compra",
    ])

    h3(doc, "Após 7 dias")
    p(doc, "Não há reembolso após o período de garantia. O acesso ao conteúdo é imediato e o valor do conhecimento entregue já pode ter sido extraído e aplicado.")

    h2(doc, "Atualizações e Modificações")
    p(doc, "A Runa Systems pode:")
    bullet(doc, [
        "Atualizar o conteúdo do programa com melhorias no método",
        "Adicionar novos módulos, documentos ou ferramentas",
        "Atualizar os agentes LENS e REEL",
        "Modificar estes termos de uso com aviso prévio por e-mail",
    ])
    p(doc, "Você será notificado por e-mail sobre atualizações importantes nos termos ou no produto.")

    h2(doc, "Cancelamento e Suspensão de Acesso")
    p(doc, "Você pode cancelar a qualquer momento. O acesso permanece ativo até o fim do período contratado.")
    p(doc, "A Runa Systems pode suspender ou encerrar seu acesso sem reembolso em caso de:")
    bullet(doc, [
        "Violação dos termos de uso",
        "Uso do programa para fins fraudulentos, antiéticos ou ilegais",
        "Compartilhamento não autorizado de acesso ou materiais",
        "Criação de conteúdo que prejudique terceiros usando o método ensinado",
        "Tentativa de reproduzir ou revender o programa sem autorização",
    ])

    h2(doc, "Jurisdição")
    p(doc, "Estes termos são regidos pelas leis brasileiras.")
    p(doc, "Foro: Comarca de Sacramento, Minas Gerais, Brasil.")

    h2(doc, "Contato")
    p(doc, "Dúvidas sobre estes termos ou sobre o programa:")
    bullet(doc, [
        "E-mail: clientesruna@gmail.com",
        "WhatsApp: 49 999527013",
    ])

    blank(doc)
    p(doc, "Ao acessar o CREATOR$, você confirma que leu, entendeu e aceita integralmente estes termos de uso — incluindo a responsabilidade pelo uso ético e legal do conhecimento adquirido.")

    path = f"{OUT_DIR}/termos-de-uso-creator-dollar.docx"
    doc.save(path)
    print(f"Generated: {path}")


if __name__ == "__main__":
    generate()
