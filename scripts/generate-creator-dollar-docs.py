"""
Generate CREATOR$ Skool course support docs as .docx files.
Uses python-docx with proper heading styles.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = "D:/Runa/runa-systems-global/SÍRIOS/📦 Entregáveis/creator-dollar-skool-docx"
os.makedirs(OUT_DIR, exist_ok=True)


def new_doc(title):
    doc = Document()
    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
    # Customize Normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    # H1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # H2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
    # H3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return doc


def h1(doc, text):
    doc.add_heading(text, level=1)

def h2(doc, text):
    doc.add_heading(text, level=2)

def h3(doc, text):
    doc.add_heading(text, level=3)

def p(doc, text):
    doc.add_paragraph(text)

def blank(doc):
    doc.add_paragraph("")

def bullet(doc, items, level=0):
    """Add bullet list items."""
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")

def code_block(doc, text):
    """Add a code/quote block with monospace font."""
    para = doc.add_paragraph()
    para.style = doc.styles["Normal"]
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    # Light grey shading
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)
    # Add left border/indent feel
    pf = OxmlElement("w:ind")
    pf.set(qn("w:left"), "360")
    pPr.append(pf)

def label_line(doc, label, value=""):
    """Bold label followed by value text."""
    para = doc.add_paragraph()
    run = para.add_run(label)
    run.bold = True
    if value:
        para.add_run(f"  {value}")


# ─────────────────────────────────────────────────────────────────────────────
# 00 — Ferramentas
# ─────────────────────────────────────────────────────────────────────────────

def doc_00():
    doc = new_doc("00-ferramentas")
    h1(doc, "As ferramentas que vamos usar")
    p(doc, "Módulo 0 · Aula 0.3")
    blank(doc)
    p(doc, "Nenhuma ferramenta é obrigatória. A lógica é transferível — se uma ferramenta mudar, o método continua o mesmo.")

    h2(doc, "Visão geral do stack")
    bullet(doc, [
        "Nano Banana Pro — geração de imagens (Passo 1) | plano gratuito: créditos diários",
        "Higgsfield — hack da pintura + input para vídeo (Passo 2) | plano gratuito",
        "Sora 2 — geração de vídeo ultra-realista (Passo 3) | acesso via Opera + VPS",
        "vmake.ai — remover marca d'água do vídeo | gratuito",
        "LENS (GPT) — otimizar prompts de imagem antes de gerar",
        "REEL (GPT) — gerar prompts narrativos para o Sora 2",
    ])

    h2(doc, "Nano Banana Pro")
    p(doc, "Gera imagens a partir de prompts de texto. Ótimo para criar a imagem base do avatar.")
    p(doc, "Alternativa: o próprio Higgsfield também gera imagens — use se preferir centralizar em uma única ferramenta.")
    p(doc, "Versão gratuita: créditos diários renovam automaticamente. Para o volume deste curso, o plano gratuito é suficiente.")

    h2(doc, "Higgsfield")
    p(doc, "Duas funções no método:")
    numbered(doc, [
        "Gerar a imagem base do avatar (alternativa ao Nano Banana Pro)",
        "Aplicar o hack da pintura ultra-realista — obrigatório antes do Sora 2",
    ])
    p(doc, "O hack do Módulo 3 — transformar a imagem em pintura com traços de tinta — só funciona no Higgsfield. Esse passo é o que garante movimento natural no vídeo.")

    h2(doc, "Sora 2")
    p(doc, "Gera vídeos de 15 segundos a partir de imagem + prompt. É a ferramenta que transforma a imagem do avatar em vídeo ultra-realista.")
    p(doc, "O Sora 2 não está disponível diretamente no Brasil. As instruções completas de acesso estão no Módulo 3.")

    h2(doc, "vmake.ai")
    p(doc, "Remove a marca d'água do Sora 2 automaticamente. Upload do vídeo, processamento automático, download do vídeo limpo. Use após selecionar o melhor take.")

    h2(doc, "Agentes incluídos no CREATOR$")

    h3(doc, "LENS — Diretor de Imagem")
    p(doc, "Otimiza o prompt antes de gerar qualquer imagem. Você descreve o avatar, ele estrutura o prompt no formato correto para o Nano Banana Pro.")
    p(doc, "Acesso: https://chatgpt.com/g/g-69be9c381dcc81919e85d56ace38c9f4-lens-r")

    h3(doc, "REEL — Diretor de Vídeo")
    p(doc, "Gera o prompt narrativo para o Sora 2. Você descreve o que quer que o avatar faça no vídeo, ele entrega o prompt completo pronto para colar.")
    p(doc, "Acesso: https://chatgpt.com/g/g-69bebdb35f748191952d6500f350c386-reel-r")

    h2(doc, "Regra de ouro")
    p(doc, "Descrição de roupa detalhada no Passo 1 = consistência no vídeo do Passo 3. O Sora 2 assume a roupa que aparece na imagem. Seja específico.")

    p(doc, "Próxima aula: Módulo 1 — Construindo a Personagem")
    doc.save(f"{OUT_DIR}/00-ferramentas.docx")
    print("✓ 00-ferramentas.docx")


# ─────────────────────────────────────────────────────────────────────────────
# 01 — Persona
# ─────────────────────────────────────────────────────────────────────────────

def doc_01():
    doc = new_doc("01-persona")
    h1(doc, "Construindo a Personagem")
    p(doc, "Módulo 1 · Aulas 1.1 e 1.2 — Identidade · Voz · Tom")
    blank(doc)
    p(doc, "Antes de abrir qualquer ferramenta de imagem, você precisa de uma definição completa da personagem. É aqui que a maioria das pessoas erra: elas tentam construir o visual antes de saber quem é o avatar.")
    p(doc, "A identidade vem antes da imagem. A imagem é a tradução visual da identidade.")

    h2(doc, "Parte 1 — Identidade da Personagem")
    p(doc, "Preencha com o máximo de detalhe. Quanto mais específico, melhor será o Prompt Mestre.")

    h3(doc, "Nome e origem")
    label_line(doc, "Nome do avatar:")
    blank(doc)
    label_line(doc, "Origem / backstory", "(de onde ela vem, o que formou quem ela é):")
    blank(doc)
    blank(doc)

    h3(doc, "Transformação e missão")
    label_line(doc, "O que ela superou", "(vulnerabilidade que a torna real, não perfeita):")
    blank(doc)
    label_line(doc, "Transformação", "(antes → depois):")
    blank(doc)
    label_line(doc, "Missão", "(o que ela quer construir no mundo):")
    blank(doc)

    h3(doc, "Valores")
    p(doc, "Liste de 3 a 5 valores centrais, em ordem de prioridade:")
    numbered(doc, ["", "", "", "", ""])

    h3(doc, "Aparência física")
    label_line(doc, "Tom de pele", "(ex: médio, oliva, claro com sardas, escuro):")
    blank(doc)
    label_line(doc, "Cabelo", "(cor, comprimento, textura):")
    blank(doc)
    label_line(doc, "Biótipo", "(ex: esguia, atlética, curvilínea, alta):")
    blank(doc)
    label_line(doc, "Característica marcante", "(opcional):")
    blank(doc)

    h3(doc, "Estilo visual")
    label_line(doc, "Paleta de cores preferida para roupa:")
    blank(doc)
    label_line(doc, "Roupa mais frequente", "(tipo + cor exata + tecido + modelagem):")
    blank(doc)
    blank(doc)
    label_line(doc, "Ambiente típico:")
    blank(doc)
    label_line(doc, "Estética geral", "(ex: iPhone casual, editorial minimalista, Lo-Fi criadora):")
    blank(doc)

    h2(doc, "Parte 2 — Voz e Tom")

    h3(doc, "Como ela se comunica")
    label_line(doc, "Vocabulário frequente", "(palavras, expressões, gírias):")
    blank(doc)
    blank(doc)
    label_line(doc, "O que ela NUNCA diz:")
    blank(doc)
    label_line(doc, "Tom emocional", "(ex: direta e confiante / acolhedora / séria e intelectual):")
    blank(doc)

    h3(doc, "Formato de conteúdo")
    label_line(doc, "Como ela inicia posts ou vídeos:")
    blank(doc)
    label_line(doc, "Comprimento típico das mensagens:")
    blank(doc)
    label_line(doc, "O que ela trata com profundidade:")
    blank(doc)
    label_line(doc, "O que ela evita ou nunca aborda:")
    blank(doc)

    h2(doc, "O que fazer com isso")
    numbered(doc, [
        "Leve para a Aula 1.3 — Preenchendo o Storyboard",
        "Use como input para o agente LENS na geração de imagens (Módulo 2)",
        "Guarde: é a base de todo conteúdo futuro do avatar",
    ])
    p(doc, "Próxima aula: 1.3 — Preenchendo o Storyboard")
    doc.save(f"{OUT_DIR}/01-persona.docx")
    print("✓ 01-persona.docx")


# ─────────────────────────────────────────────────────────────────────────────
# 02 — Storyboard Template
# ─────────────────────────────────────────────────────────────────────────────

def doc_02():
    doc = new_doc("02-storyboard-template")
    h1(doc, "Storyboard — Template Completo")
    p(doc, "Módulo 1 · Aula 1.3 — O exato template usado para construir o Alpha®")
    blank(doc)
    p(doc, "O Storyboard é o documento permanente que torna o seu avatar reproduzível a qualquer momento, em qualquer ferramenta, por qualquer pessoa — incluindo agentes de IA.")
    p(doc, "Ele é composto por 6 blocos:")
    numbered(doc, [
        "Identidade",
        "Descrição visual completa",
        "Prompt Mestre validado",
        "Biblioteca de imagens",
        "Prompt de vídeo validado",
        "Voz e tom",
    ])

    h2(doc, "Bloco 1 — Identidade")
    label_line(doc, "Nome:")
    blank(doc)
    label_line(doc, "Tagline", "(uma frase que define quem ela é):")
    blank(doc)
    label_line(doc, "Origem e backstory:")
    blank(doc)
    blank(doc)
    label_line(doc, "Transformação", "(de onde veio / onde está / para onde vai):")
    blank(doc)
    blank(doc)
    label_line(doc, "Vulnerabilidade", "(o que a torna humana):")
    blank(doc)
    label_line(doc, "Missão:")
    blank(doc)
    label_line(doc, "Valores", "(3 a 5, em ordem de prioridade):")
    numbered(doc, ["", "", "", "", ""])

    h2(doc, "Bloco 2 — Descrição Visual Completa")

    h3(doc, "Aparência")
    for lbl in ["Tom de pele:", "Olhos (cor, formato):", "Cabelo (cor, comprimento, textura):", "Biótipo:", "Característica marcante:"]:
        label_line(doc, lbl)
        blank(doc)

    h3(doc, "Look Principal")
    p(doc, "Descreva a roupa que ela usa com mais frequência, peça a peça:")
    label_line(doc, "Peça 1 — Tipo:")
    label_line(doc, "Peça 1 — Cor exata:")
    label_line(doc, "Peça 1 — Tecido:")
    label_line(doc, "Peça 1 — Modelagem:")
    blank(doc)
    label_line(doc, "Peça 2 — Tipo:")
    label_line(doc, "Peça 2 — Cor exata:")
    label_line(doc, "Peça 2 — Tecido:")
    label_line(doc, "Peça 2 — Modelagem:")
    blank(doc)
    label_line(doc, "Calçado (se relevante):")

    h3(doc, "Looks Alternativos")
    label_line(doc, "Look 2 — Nome:")
    blank(doc)
    label_line(doc, "Look 2 — Descrição:")
    blank(doc)
    blank(doc)
    label_line(doc, "Look 3 — Nome:")
    blank(doc)
    label_line(doc, "Look 3 — Descrição:")
    blank(doc)

    h3(doc, "Paleta de Cores")
    for lbl in ["Cor primária:", "Cor secundária:", "Cor de acento:", "Cor que NUNCA aparece:"]:
        label_line(doc, lbl)

    h3(doc, "Ambientes")
    label_line(doc, "Ambiente principal:")
    blank(doc)
    p(doc, "Ambientes secundários:")
    numbered(doc, ["", "", ""])
    label_line(doc, "Iluminação preferida:")
    blank(doc)

    h2(doc, "Bloco 3 — Prompt Mestre")
    label_line(doc, "Prompt principal", "(em inglês — colar direto nas ferramentas):")
    blank(doc)
    blank(doc)
    blank(doc)
    label_line(doc, "Prompt negativo:")
    blank(doc)
    blank(doc)
    h3(doc, "Variações aprovadas")
    for v in ["Look casual:", "Look profissional:", "Ambiente externo:", "Close-up / expressão:"]:
        label_line(doc, v)

    h2(doc, "Bloco 4 — Biblioteca de Imagens")
    p(doc, "Estrutura de pasta:")
    code_block(doc,
        "avatar-[nome]/\n"
        "  imagens/\n"
        "    referencia-base.png      ← imagem-base original (Passo 1)\n"
        "    pintura-higgsfield.png   ← imagem-pintura (Passo 2)\n"
        "    curadas/                 ← 20-30 melhores gerações\n"
        "  videos/\n"
        "    curados/                 ← 3-5 melhores takes"
    )
    blank(doc)
    p(doc, "Critério de seleção:")
    bullet(doc, [
        "Consistência facial entre gerações",
        "Qualidade de iluminação natural",
        "Postura e expressão naturais (não posada)",
    ])

    h2(doc, "Bloco 5 — Prompt de Vídeo")
    label_line(doc, "Prompt Sora 2", "(gerado pelo agente REEL):")
    blank(doc)
    blank(doc)
    h3(doc, "Configuração usada")
    bullet(doc, [
        "Style: Selfie",
        "Duração: 15 segundos",
        "Input: imagem-pintura do Passo 2",
        "Takes gerados: 3 — take selecionado: ____",
    ])
    p(doc, "Critério de seleção:")
    bullet(doc, [
        "Movimento fluido (sem artefatos)",
        "Rosto consistente com as imagens curadas",
        "Naturalidade — parece gravado, não gerado",
    ])

    h2(doc, "Bloco 6 — Voz e Tom")
    label_line(doc, "Palavras que ela usa com frequência:")
    blank(doc)
    blank(doc)
    label_line(doc, "Palavras que ela NUNCA usa:")
    blank(doc)
    label_line(doc, "Tom emocional:")
    blank(doc)
    label_line(doc, "Como inicia posts e reels:")
    blank(doc)
    label_line(doc, "Estrutura típica de copy:")
    blank(doc)
    label_line(doc, "Exemplo de copy 1:")
    blank(doc)
    blank(doc)
    label_line(doc, "Exemplo de copy 2:")
    blank(doc)
    blank(doc)

    h2(doc, "Como usar este Storyboard")
    bullet(doc, [
        "Preencha bloco por bloco ao longo do programa",
        "Revise o Bloco 3 sempre que validar novas imagens",
        "Compartilhe com agentes de IA, designers ou colaboradores para replicar o avatar sem briefings verbais",
        "No MIND$, este documento se torna o input do sistema de prompts da personagem",
    ])
    doc.save(f"{OUT_DIR}/02-storyboard-template.docx")
    print("✓ 02-storyboard-template.docx")


# ─────────────────────────────────────────────────────────────────────────────
# 03 — Master Prompt Guide
# ─────────────────────────────────────────────────────────────────────────────

def doc_03():
    doc = new_doc("03-master-prompt-guide")
    h1(doc, "Guia do Prompt Mestre")
    p(doc, "Módulo 2 · Aula 2.2")
    blank(doc)
    p(doc, "O Prompt Mestre é o bloco que garante que a mesma personagem aparece em todas as gerações. É o documento mais importante do método — sem ele, cada imagem parece uma pessoa diferente.")

    h2(doc, "O problema que ele resolve")
    p(doc, "Ferramentas de IA generativa não têm memória entre sessões. Cada vez que você gera uma imagem, é como se a ferramenta nunca tivesse visto o seu avatar antes.")
    p(doc, "O Prompt Mestre resolve isso descrevendo a personagem em tanto detalhe que a ferramenta não precisa lembrar — ela reconstrói a identidade a partir da descrição.")

    h2(doc, "Estrutura do Prompt Mestre")
    p(doc, "Um prompt eficaz tem 6 blocos:")
    numbered(doc, [
        "Contexto da cena — o que está acontecendo, de onde a foto está sendo tirada",
        "Personagem — descrição física: tom de pele, cabelo, biótipo, característica marcante",
        "Roupa — peça a peça: tipo, cor exata, tecido, modelagem, caimento no corpo",
        "Ambiente — onde ela está, o que está ao redor, perspectiva",
        "Iluminação — fonte de luz, temperatura, sombras, estilo",
        "Estilo fotográfico — tipo de câmera simulada, composição, saturação, vibe geral",
    ])
    p(doc, "NEGATIVE: lista do que deve ser excluído")

    h2(doc, "Exemplo completo — Alpha®")
    code_block(doc,
        "Casual indoor photo of this woman from the reference photos,\n"
        "speaking in the middle of a sentence, slightly warmer and\n"
        "healthier skin tone, with a subtle glow (no retouching),\n"
        "one hand naturally raised near her chest in a relaxed gesture.\n\n"
        "She is wearing:\n"
        "- a white ribbed sleeveless tank top, cropped at the waist\n"
        "- high neckline, snug but natural fit\n"
        "- visible ribbed texture with realistic fabric tension and folds\n"
        "- light beige / off-white loose pants\n"
        "- elastic waistband with a visible drawstring tied at the front\n"
        "- relaxed, casual fit with authentic fabric behavior and movement\n\n"
        "Soft ambient lighting in an indoor environment,\n"
        "warm lights from room decor, soft shadows that\n"
        "define her cheekbones and jawline.\n\n"
        "Neutral wall and casual room elements behind her.\n"
        "Photo casually taken with an iPhone by a friend, slightly off-center.\n"
        "Natural expression, no posing, authentic moment, soft yet warm color palette.\n\n"
        "NEGATIVE:\n"
        "professional photography, editorial studio lighting,\n"
        "golden hour, intense bokeh, portrait mode,\n"
        "HDR, oversaturated colors, white clothes, pale skin tone, cool tones"
    )

    h2(doc, "Regras críticas")

    h3(doc, "Roupa em detalhes — sempre")
    p(doc, "A roupa é o elemento de consistência mais importante. Descreva: tipo da peça, cor exata, tecido, como cai no corpo.")
    p(doc, "O Sora 2 assume a roupa da imagem-base. Roupa inconsistente nas imagens = vídeo inconsistente.")

    h3(doc, "Iluminação Lo-Fi — nunca estúdio")
    p(doc, "Use: luz natural de janela, luz ambiente de sala (warm lights from room decor), luz difusa ao ar livre.")
    p(doc, "Evite no prompt e no bloco NEGATIVE: studio lighting, professional photography, editorial, golden hour, HDR.")

    h3(doc, "Expressão autêntica — nunca posada")
    p(doc, "Descreva uma ação natural em andamento:")
    bullet(doc, [
        '"speaking in the middle of a sentence"',
        '"looking down at her phone"',
        '"laughing at something just heard"',
        '"one hand naturally raised near her chest"',
    ])

    h2(doc, "Como usar o agente LENS")
    p(doc, "O LENS é um GPT treinado para otimizar prompts no estilo Lo-Fi realista. Use-o antes de gerar qualquer imagem.")
    p(doc, "Acesso: https://chatgpt.com/g/g-69be9c381dcc81919e85d56ace38c9f4-lens-r")
    numbered(doc, [
        "Cole a descrição completa do seu avatar (aparência, roupa, ambiente, tom)",
        'Peça: "Gere um prompt otimizado para Nano Banana Pro no estilo Lo-Fi realista"',
        "Use o prompt gerado diretamente na ferramenta",
    ])

    h2(doc, "Checklist antes de gerar")
    bullet(doc, [
        "Roupa descrita com tipo, cor exata, tecido e modelagem",
        "Ambiente definido (interior / exterior, elementos de fundo)",
        "Iluminação especificada (não deixar vago)",
        "Expressão descrita como ação natural em andamento",
        "Bloco NEGATIVE inclui: studio lighting, HDR, portrait mode",
        "Fundo neutro selecionado",
    ])

    h2(doc, "Corrigindo o prompt por problema")
    bullet(doc, [
        "Pele muito clara → adicione \"warm skin tone, healthy glow\"",
        "Iluminação fria → substitua por \"warm ambient light from room decor\"",
        "Parece foto de estúdio → adicione \"iPhone photo by a friend, slightly off-center\"",
        "Roupa diferente do descrito → mova a roupa para mais cedo no prompt; detalhe o tecido",
        "Fundo elaborado → adicione \"simple neutral background, minimal room elements\"",
    ])
    p(doc, "Próxima aula: 2.3–2.5 — Gerando e Curando Imagens")
    doc.save(f"{OUT_DIR}/03-master-prompt-guide.docx")
    print("✓ 03-master-prompt-guide.docx")


# ─────────────────────────────────────────────────────────────────────────────
# 04 — Biblioteca Visual
# ─────────────────────────────────────────────────────────────────────────────

def doc_04():
    doc = new_doc("04-biblioteca-visual")
    h1(doc, "Biblioteca Visual")
    p(doc, "Módulo 2 · Aulas 2.3 a 2.5 — Geração · Variações · Curadoria")
    blank(doc)
    p(doc, "A biblioteca visual é o acervo de imagens aprovadas do avatar. É a base de todo conteúdo futuro — posts, thumbnails, reels.")
    p(doc, "Meta ao final do Módulo 2: 30+ imagens curadas e organizadas.")

    h2(doc, "Travando o estilo — a imagem-referência")
    p(doc, "Antes de gerar variações, você precisa de uma imagem-base aprovada — a referência mestra que define o padrão para todas as outras.")
    numbered(doc, [
        "Gere 5 imagens com o Prompt Mestre",
        "Selecione a melhor pelo critério abaixo",
        "Salve como referencia-base.png",
        "Use como referência visual nas próximas gerações",
    ])
    h3(doc, "Critério de seleção da referência-base")
    bullet(doc, [
        "O rosto corresponde à descrição da personagem?",
        "A luz é natural, não de estúdio?",
        "A roupa está exatamente como descrita no prompt?",
        "A expressão parece natural, não posada?",
        "O fundo é neutro o suficiente?",
    ])

    h2(doc, "Gerando variações")
    p(doc, "Modifique um elemento por vez. O Prompt Mestre base permanece em tudo mais.")

    h3(doc, "Variações de expressão")
    p(doc, "Adicione ao final do Prompt Mestre:")
    bullet(doc, [
        '"laughing softly at something just said"',
        '"looking directly at camera with a slight smirk"',
        '"looking down, distracted, caught off guard"',
        '"talking, mid-sentence, animated"',
    ])

    h3(doc, "Variações de postura")
    bullet(doc, [
        '"both hands holding a coffee mug"',
        '"one hand running through hair"',
        '"arms crossed lightly, relaxed"',
        '"leaning slightly against the wall"',
    ])

    h3(doc, "Variações de look")
    p(doc, "Para cada look alternativo: mantenha contexto + personagem + ambiente + estilo. Substitua apenas o bloco de roupa pelo look novo.")

    h3(doc, "Variações de ambiente")
    p(doc, "Substitua o bloco de ambiente por:")
    bullet(doc, [
        '"seated at a wooden desk with warm light from a lamp"',
        '"standing near a window with natural daylight, soft curtain diffusion"',
        '"in a coffee shop, brick walls, warm ambient light"',
    ])

    h2(doc, "Seeds e referência de imagem")
    p(doc, "O seed fixa o ponto de partida da geração. Mesmo seed com pequenas variações no prompt tende a manter maior consistência facial.")
    bullet(doc, [
        "Anote o seed da referencia-base.png (a ferramenta exibe após geração)",
        "Use o mesmo seed para variações de expressão e postura",
        "Para variações de look ou ambiente, o seed é menos crítico",
    ])
    p(doc, "Image-to-Image: no Nano Banana Pro e Higgsfield, anexe a referencia-base.png como referência visual. A ferramenta mantém os traços faciais enquanto aplica as variações do prompt.")

    h2(doc, "Selecionando e curando")
    h3(doc, "Quantas imagens gerar por sessão")
    p(doc, "Gere em lotes de 5. De 5 imagens, espere aprovar 2 a 3.")
    bullet(doc, [
        "Expressões: 3 lotes → meta de 6–8 aprovadas",
        "Posturas: 2 lotes → meta de 4–5 aprovadas",
        "Looks alternativos: 2 lotes por look → meta de 4–5 por look",
        "Ambientes: 2 lotes → meta de 4–5 aprovadas",
        "Total esperado: 30–40 geradas → 20–30 aprovadas",
    ])

    h3(doc, "Checklist de curadoria por imagem")
    bullet(doc, [
        "Rosto consistente com a referencia-base.png",
        "Roupa corresponde ao look descrito",
        "Iluminação natural (não de estúdio)",
        "Expressão espontânea, não posada",
        "Resolução adequada para post no Instagram",
        "Fundo não distrai da personagem",
    ])

    h3(doc, "O que descartar imediatamente")
    bullet(doc, [
        "Artefatos visíveis (mãos deformadas, objetos multiplicados)",
        "Texto ou letras geradas pela IA",
        "Expressão que parece rendering 3D ou boneco",
        "Roupa diferente do descrito no prompt",
        "Rosto com traços que contradizem a definição da personagem",
    ])

    h2(doc, "Organizando a biblioteca")
    code_block(doc,
        "avatar-[nome]/\n"
        "  imagens/\n"
        "    referencia-base.png\n"
        "    pintura-higgsfield.png\n"
        "    curadas/\n"
        "      look-01-casual/\n"
        "      look-02-[nome]/\n"
        "      expressoes/\n"
        "  videos/\n"
        "    curados/"
    )
    p(doc, "Próxima aula: Módulo 3 — O Hack do Sora 2")
    doc.save(f"{OUT_DIR}/04-biblioteca-visual.docx")
    print("✓ 04-biblioteca-visual.docx")


# ─────────────────────────────────────────────────────────────────────────────
# 05 — Hack Sora
# ─────────────────────────────────────────────────────────────────────────────

def doc_05():
    doc = new_doc("05-hack-sora")
    h1(doc, "O Hack do Sora 2")
    p(doc, "Módulo 3 · Aula 3.2 — Higgsfield → Pintura → Vídeo Ultra-Realista")
    blank(doc)
    p(doc, "Este é o passo que faz a diferença entre um vídeo com rosto inconsistente e um vídeo com movimento fluido e identidade preservada. Não pule.")

    h2(doc, "Por que isso funciona")
    p(doc, "O Sora 2 é excelente em gerar movimento natural — mas quando recebe uma foto realista como input, ele tende a corrigir o rosto baseado nos padrões que aprendeu. O resultado: o rosto muda.")
    p(doc, "A solução é transformar a foto em uma pintura com traços de tinta visíveis antes de enviar ao Sora 2. O modelo interpreta a pintura como referência artística — não como foto real. Isso gera movimento mais natural, mais fluido e com o rosto preservado.")

    h2(doc, "Passo 1 — Preparar a imagem-base")
    p(doc, "Use a imagem aprovada do Módulo 2. Idealmente:")
    bullet(doc, [
        "Fundo neutro (facilita o Sora a focar no movimento da personagem)",
        "Rosto desobstruído (sem óculos escuros cobrindo a face)",
        "Roupa completa e visível (o Sora assume a roupa da imagem)",
    ])

    h2(doc, "Passo 2 — Aplicar o hack no Higgsfield")
    numbered(doc, [
        "Abra o Higgsfield (higgsfield.ai)",
        "Carregue a imagem do Passo 1",
        "Use exatamente o prompt abaixo",
        "Gere e salve como pintura-higgsfield.png",
    ])
    code_block(doc,
        "transforme essa imagem em uma pintura ultra realista\n"
        "com todas as características da personagem,\n"
        "mas traga marcas de tinta visíveis"
    )
    blank(doc)
    p(doc, "Não modifique este prompt. A combinação exata \"pintura ultra realista\" + \"marcas de tinta visíveis\" é o que ativa o comportamento correto no Sora 2.")

    h2(doc, "Passo 3 — Gerar o vídeo no Sora 2")

    h3(doc, "Acesso no Brasil")
    p(doc, "O Sora 2 não está disponível diretamente no Brasil. As instruções completas de acesso estão no Módulo 3 em vídeo.")

    h3(doc, "Usando o agente REEL")
    p(doc, "Antes de abrir o Sora 2, gere o prompt de vídeo com o agente REEL.")
    p(doc, "Acesso: https://chatgpt.com/g/g-69bebdb35f748191952d6500f350c386-reel-r")
    numbered(doc, [
        'Descreva o que você quer que o avatar faça no vídeo (ex: "ela está falando diretamente para a câmera, natural, como numa story")',
        '"Gere um prompt para Sora 2 estilo selfie natural, 15 segundos"',
        "Use o prompt gerado no campo de texto do Sora 2",
    ])

    h3(doc, "Configuração no Sora 2")
    bullet(doc, [
        "Input: pintura-higgsfield.png (não a foto original)",
        "Prompt: output do agente REEL",
        "Style: Selfie",
        "Duração: 15 segundos",
        "Repita 3 vezes com o mesmo prompt e a mesma imagem",
    ])

    h2(doc, "Passo 4 — Selecionar o melhor take")
    h3(doc, "Aprovar se")
    bullet(doc, [
        "Movimento fluido — sem artefatos, sem teleportes",
        "Rosto consistente com as imagens curadas",
        "Roupa não mudou no meio do vídeo",
        "Naturalidade — parece gravado, não gerado",
    ])
    h3(doc, "Descartar se")
    bullet(doc, [
        "Olhos piscando de forma estranha ou assimétrica",
        "Boca se movendo sem lógica (se o prompt não era de fala)",
        "Fundo com elementos que aparecem e desaparecem",
        "Deformação de mãos ou pescoço",
    ])

    h2(doc, "Passo 5 — Remover marca d'água")
    numbered(doc, [
        "Acesse vmake.ai",
        "Faça upload do take selecionado",
        "Processamento automático",
        "Baixe o vídeo limpo",
    ])
    p(doc, "Salve como take-01.mp4 em avatar-[nome]/videos/curados/")

    h2(doc, "Resumo do fluxo")
    code_block(doc,
        "referencia-base.png  (Módulo 2)\n"
        "       ↓\n"
        "Higgsfield → \"transforme em pintura ultra realista com marcas de tinta\"\n"
        "       ↓\n"
        "pintura-higgsfield.png\n"
        "       ↓\n"
        "Agente REEL → prompt de vídeo\n"
        "       ↓\n"
        "Sora 2  (Style: Selfie · 15s · 3x gerações)\n"
        "       ↓\n"
        "Selecionar melhor take\n"
        "       ↓\n"
        "vmake.ai (remover marca d'água)\n"
        "       ↓\n"
        "take-01.mp4 → pronto para publicar"
    )
    p(doc, "Próxima aula: 3.4 — Biblioteca de Formatos")
    doc.save(f"{OUT_DIR}/05-hack-sora.docx")
    print("✓ 05-hack-sora.docx")


# ─────────────────────────────────────────────────────────────────────────────
# 06 — Formatos de Exportação
# ─────────────────────────────────────────────────────────────────────────────

def doc_06():
    doc = new_doc("06-formatos-exportacao")
    h1(doc, "Biblioteca de Formatos — Exportação por Destino")
    p(doc, "Módulo 3 · Aula 3.4")
    blank(doc)
    p(doc, "O mesmo vídeo ou imagem gerado pode ser adaptado para múltiplos formatos. Use esta referência sempre que for exportar.")

    h2(doc, "Formatos de vídeo")

    h3(doc, "Reels do Instagram — formato principal")
    bullet(doc, [
        "Proporção: 9:16 (vertical)",
        "Resolução: 1080 × 1920 px",
        "Duração: 7 a 90 segundos (ideal: 15–30s)",
        "FPS: 30",
        "Formato: MP4 (H.264)",
        "Áudio: AAC, 44.1 kHz",
        "Tamanho máximo: 250 MB",
    ])
    p(doc, "Nota: o Sora 2 gera em 16:9 por padrão. Para Reels, recorte verticalmente no CapCut centralizando o rosto.")

    h3(doc, "Stories do Instagram")
    bullet(doc, [
        "Proporção: 9:16 (vertical)",
        "Resolução: 1080 × 1920 px",
        "Duração: até 60 segundos por segmento",
        "Formato: MP4 (H.264)",
    ])
    p(doc, "Zona segura: mantenha o rosto e elementos principais entre y=250px e y=1670px para não ser cortado pelo UI do Instagram.")

    h3(doc, "Feed do Instagram — vídeo quadrado")
    bullet(doc, ["Proporção: 1:1", "Resolução: 1080 × 1080 px", "Duração: até 60 segundos", "Formato: MP4 (H.264)"])

    h3(doc, "YouTube Shorts")
    bullet(doc, ["Proporção: 9:16 (vertical)", "Resolução: 1080 × 1920 px", "Duração: até 60 segundos", "Formato: MP4"])

    h2(doc, "Formatos de imagem")

    h3(doc, "Post de feed — quadrado")
    bullet(doc, ["Proporção: 1:1", "Resolução: 1080 × 1080 px", "Formato: JPG ou PNG", "Qualidade JPG: 80–90%"])

    h3(doc, "Post de feed — retrato (melhor cobertura no feed)")
    bullet(doc, ["Proporção: 4:5", "Resolução: 1080 × 1350 px", "Formato: JPG ou PNG"])

    h3(doc, "Carrossel")
    bullet(doc, [
        "Proporção: 1:1 ou 4:5 (manter consistente entre slides)",
        "Resolução: 1080 × 1080 px ou 1080 × 1350 px",
        "Slides: 2 a 10",
        "Formato: JPG ou PNG por slide",
    ])

    h3(doc, "Thumbnail YouTube")
    bullet(doc, [
        "Proporção: 16:9",
        "Resolução mínima: 1280 × 720 px",
        "Resolução recomendada: 1920 × 1080 px",
        "Formato: JPG ou PNG",
        "Tamanho máximo: 2 MB",
    ])

    h2(doc, "Recorte 16:9 → 9:16 no CapCut")
    numbered(doc, [
        "Importe o vídeo no CapCut",
        "Selecione Proporção → 9:16",
        "Reposicione para centralizar o rosto",
        "Verifique que o rosto não é cortado em nenhum momento",
        "Exporte em 1080 × 1920, 30fps",
    ])

    h2(doc, "Checklist antes de publicar")
    bullet(doc, [
        "Proporção correta para o formato",
        "Resolução ≥ 1080px na menor dimensão",
        "Rosto visível e centralizado (sem corte pelo UI)",
        "Marca d'água removida (vmake.ai)",
        "Áudio adicionado se necessário",
        "Legenda ou texto não cobre o rosto",
    ])
    p(doc, "Próxima aula: Módulo 4 — Montando o Storybook")
    doc.save(f"{OUT_DIR}/06-formatos-exportacao.docx")
    print("✓ 06-formatos-exportacao.docx")


# ─────────────────────────────────────────────────────────────────────────────
# 07 — Storybook Final
# ─────────────────────────────────────────────────────────────────────────────

def doc_07():
    doc = new_doc("07-storybook-final")
    h1(doc, "Montando o Storybook Final")
    p(doc, "Módulo 4 · Aula 4.2")
    blank(doc)
    p(doc, "O Storybook é o documento definitivo do seu avatar. Ele transforma tudo que foi construído nos módulos anteriores em uma referência permanente — usável por IA, por humanos, por equipes inteiras.")

    h2(doc, "O que você já tem neste ponto")
    bullet(doc, [
        "Identidade da personagem (Módulo 1 — Aulas 1.1 e 1.2)",
        "Voz e tom (Módulo 1 — Aula 1.2)",
        "Prompt Mestre validado (Módulo 2 — Aula 2.2)",
        "Biblioteca de imagens (Módulo 2 — Aulas 2.3 a 2.5)",
        "Prompt de vídeo validado (Módulo 3 — Aula 3.2)",
        "Formatos aprovados (Módulo 3 — Aula 3.4)",
    ])
    p(doc, "A montagem é consolidar tudo isso em um único documento. Se você foi preenchendo o template ao longo do programa, a maior parte já está lá.")

    h2(doc, "Onde montar")
    bullet(doc, [
        "Notion — ideal para banco de imagens embutido, fácil de compartilhar via link",
        "Google Docs — mais simples, mais universal, ótimo para colaboradores externos",
    ])

    h2(doc, "Revisão bloco a bloco")

    h3(doc, "Bloco 1 — Identidade")
    bullet(doc, [
        "Nome e tagline definidos",
        "Origem / backstory em 3 a 5 frases",
        "Transformação: antes → depois → futuro",
        "Vulnerabilidade documentada",
        "Missão em 1 frase",
        "3 a 5 valores em ordem de prioridade",
    ])
    p(doc, "Teste: alguém que nunca ouviu falar do avatar consegue descrevê-la após ler este bloco?")

    h3(doc, "Bloco 2 — Descrição Visual")
    bullet(doc, [
        "Aparência física completa",
        "Look principal descrito peça a peça",
        "2 looks alternativos definidos e nomeados",
        "Paleta de cores documentada",
        "3 ambientes definidos",
        "Estilo de iluminação especificado",
    ])
    p(doc, "Teste: um designer consegue gerar a personagem sem ver nenhuma foto — só com este bloco?")

    h3(doc, "Bloco 3 — Prompt Mestre")
    bullet(doc, [
        "Prompt principal em inglês (mínimo 6 linhas)",
        "Bloco NEGATIVE com 5 a 8 elementos",
        "Tabela de variações aprovadas",
        "Seed da referência-base anotado (se disponível)",
    ])

    h3(doc, "Bloco 4 — Biblioteca de Imagens")
    bullet(doc, [
        "referencia-base.png identificada",
        "pintura-higgsfield.png salva",
        "20+ imagens curadas organizadas por look",
    ])

    h3(doc, "Bloco 5 — Prompt de Vídeo")
    bullet(doc, [
        "Prompt do REEL documentado",
        "Configurações do Sora 2 registradas",
        "3 a 5 vídeos curados salvos",
    ])

    h3(doc, "Bloco 6 — Voz e Tom")
    bullet(doc, [
        "Vocabulário frequente listado",
        "Palavras proibidas listadas",
        "Tom emocional descrito",
        "Estrutura típica de copy definida",
        "2 exemplos de copy aprovados incluídos",
    ])
    p(doc, "Teste: um copywriter externo consegue escrever como o avatar após ler este bloco?")

    h2(doc, "Como usar o Storybook pronto")
    h3(doc, "Para você mesmo")
    bullet(doc, [
        "Abra o Bloco 3 antes de gerar qualquer imagem nova",
        "Não tente lembrar de memória — use o documento como referência",
    ])
    h3(doc, "Para agentes de IA")
    bullet(doc, [
        'Cole o Storyboard completo no contexto do agente',
        '"Este é o documento de referência de [Nome]. Use-o para orientar todas as gerações e textos."',
    ])
    h3(doc, "Para colaboradores humanos")
    bullet(doc, [
        "Compartilhe o link (Notion / Google Docs)",
        "Não precisa de briefing verbal — o Storybook faz isso",
    ])
    h3(doc, "Para os próximos produtos")
    bullet(doc, [
        "MIND$: Bloco 6 + Bloco 1 = input do sistema de prompts do avatar",
        "$QUAD: Storybook completo = briefing da squad de conteúdo",
    ])

    h2(doc, "Sinais de que o Storybook está pronto")
    bullet(doc, [
        "Você gera uma nova imagem sem abrir o histórico de gerações",
        "Alguém externo leu e conseguiu descrever o avatar corretamente",
        "O Prompt Mestre gera resultados consistentes em 3 de 5 gerações",
        "Você tem pelo menos 20 imagens e 3 vídeos aprovados na biblioteca",
    ])
    doc.save(f"{OUT_DIR}/07-storybook-final.docx")
    print("✓ 07-storybook-final.docx")


# ─────────────────────────────────────────────────────────────────────────────
# 08 — Prompt Pack
# ─────────────────────────────────────────────────────────────────────────────

def doc_08():
    doc = new_doc("08-prompt-pack")
    h1(doc, "Prompt Pack — 20+ Prompts Validados")
    p(doc, "Bônus — Os prompts usados para construir o Alpha®")
    blank(doc)
    p(doc, "Organizados por categoria. Substitua os placeholders [ENTRE COLCHETES] pelos dados do seu avatar.")
    bullet(doc, [
        "[TOM DE PELE] → ex: olive-toned, warm brown, light with freckles",
        "[ROUPA 1] → ex: white ribbed tank top, cropped at the waist",
        "[ROUPA 2] → ex: light beige linen pants, relaxed fit",
    ])

    h2(doc, "Categoria 1 — Retratos (Portraits)")

    h3(doc, "P-01 — Casual indoor · falando para câmera")
    code_block(doc,
        "Casual indoor photo of a [TOM DE PELE] woman,\n"
        "speaking directly to camera in the middle of a sentence,\n"
        "warm and healthy skin tone with a subtle natural glow,\n"
        "one hand lightly raised near chest in a relaxed gesture.\n\n"
        "She is wearing:\n- [ROUPA 1]\n- [ROUPA 2]\n\n"
        "Soft ambient lighting from warm room decor,\n"
        "soft shadows defining cheekbones and jawline.\n"
        "Neutral wall behind her, minimal casual room elements.\n"
        "iPhone photo casually taken by a friend, slightly off-center.\n"
        "Natural expression, authentic moment.\n\n"
        "NEGATIVE: professional photography, studio lighting, golden hour,\n"
        "intense bokeh, portrait mode, HDR, oversaturated, editorial"
    )

    h3(doc, "P-02 — Close-up · expressão pensativa")
    code_block(doc,
        "Close-up iPhone photo of a [TOM DE PELE] woman,\n"
        "slightly looking down and to the side, thoughtful expression,\n"
        "caught in a quiet moment, natural and unposed.\n\n"
        "[DESCREVA ROUPA COMPLETA]\n\n"
        "Diffused natural window light falling softly on her face,\n"
        "warm interior background, out of focus. Slight grain.\n\n"
        "NEGATIVE: studio lighting, direct flash, sharp background,\n"
        "posed smile, professional retouching, HDR"
    )

    h3(doc, "P-03 — Sentada · ambiente aconchegante")
    code_block(doc,
        "Candid photo of a [TOM DE PELE] woman sitting casually,\n"
        "leaning slightly forward, relaxed posture, natural smile.\n\n"
        "[DESCREVA ROUPA COMPLETA]\n\n"
        "Cozy interior: wooden surface, warm lamp light, soft shadows.\n"
        "Shot from slightly above eye level, iPhone casual framing.\n\n"
        "NEGATIVE: restaurant setting, formal posture, professional lighting,\n"
        "cold tones, oversaturated, mirror reflections"
    )

    h3(doc, "P-04 — Em movimento · de lado")
    code_block(doc,
        "Natural candid of a [TOM DE PELE] woman walking,\n"
        "caught mid-step, looking slightly ahead,\n"
        "relaxed arms, authentic movement, not modeling.\n\n"
        "[DESCREVA ROUPA COMPLETA]\n\n"
        "Soft natural daylight, warm tones,\n"
        "slight motion blur suggesting movement.\n"
        "iPhone street photography style.\n\n"
        "NEGATIVE: runway walk, model posing, direct camera look,\n"
        "studio background, dramatic lighting"
    )

    h3(doc, "P-05 — Ação cotidiana · café")
    code_block(doc,
        "Lifestyle iPhone photo of a [TOM DE PELE] woman\n"
        "holding a coffee mug with both hands, looking at it,\n"
        "steam visible, warm morning light, quiet moment.\n\n"
        "[DESCREVA ROUPA COMPLETA]\n\n"
        "Kitchen or living room background, soft warm light.\n"
        "Natural, morning energy, unposed.\n\n"
        "NEGATIVE: stock photo, perfect posture, direct eye contact,\n"
        "magazine framing, cold lighting, artificial background"
    )

    h2(doc, "Categoria 2 — Cenários (Scenarios)")

    h3(doc, "S-01 — Ambiente de trabalho")
    code_block(doc,
        "Over-shoulder view of a [TOM DE PELE] woman\n"
        "sitting at a wooden desk, working on a laptop,\n"
        "screen glow on her face, relaxed concentration.\n\n"
        "[DESCREVA ROUPA COMPLETA]\n\n"
        "Warm desk lamp, plant or books in background,\n"
        "cozy home office, early evening warm tones.\n"
        "iPhone photo from slightly behind and to the side.\n\n"
        "NEGATIVE: corporate office, fluorescent lighting,\n"
        "formal attire, direct camera look, posed"
    )

    h3(doc, "S-02 — Ambiente externo · luz natural")
    code_block(doc,
        "Outdoor candid of a [TOM DE PELE] woman\n"
        "standing near a window or wall, indirect sunlight,\n"
        "looking slightly off-camera, relaxed.\n\n"
        "[DESCREVA ROUPA COMPLETA]\n\n"
        "Soft morning or afternoon light, soft shadows,\n"
        "background soft and out of focus.\n"
        "iPhone street style photography.\n\n"
        "NEGATIVE: harsh sunlight, fashion shoot framing,\n"
        "heavy shadow on face, bright midday sun, HDR"
    )

    h3(doc, "S-03 — Café · urbano")
    code_block(doc,
        "Candid photo of a [TOM DE PELE] woman in a café,\n"
        "seated, looking out the window, casual and natural.\n\n"
        "[DESCREVA ROUPA COMPLETA]\n\n"
        "Warm café interior, brick walls or wooden elements,\n"
        "ambient light from windows and warm bulbs.\n"
        "iPhone, slightly underexposed, authentic.\n\n"
        "NEGATIVE: posed, looking at camera, professional photography,\n"
        "coffee advertisement styling"
    )

    h2(doc, "Categoria 3 — Expressões (adicionar ao base)")
    p(doc, "Adicione ao final do seu Prompt Mestre base:")
    bullet(doc, [
        'E-01 Sorriso genuíno: "laughing softly at something just heard, eyes slightly squinting, caught mid-laugh"',
        'E-02 Direta e confiante: "looking directly at camera with calm confidence, slight smirk, knowing expression"',
        'E-03 Concentrada: "eyes cast slightly down, reading or thinking, focused expression, slight tension in brow"',
        'E-04 Surpresa positiva: "eyes slightly wide, soft surprised smile, eyebrows naturally raised"',
        'E-05 Falando / ensinando: "mid-sentence, one hand gesturing expressively near shoulder, animated"',
    ])

    h2(doc, "Categoria 4 — Variações de Look")

    h3(doc, "L-01 — Swap de look")
    p(doc, "Mantenha: abertura + personagem + ambiente + iluminação + estilo. Substitua apenas o bloco de roupa:")
    code_block(doc, "She is wearing:\n- [NOVA PEÇA 1]\n- [NOVA PEÇA 2]")

    h3(doc, "L-02 — Versão noturna / evento")
    p(doc, "Base do seu avatar + modificações:")
    bullet(doc, [
        "Roupa: [descreva look mais elaborado]",
        'Ambiente: "warm restaurant interior, candlelight, evening atmosphere"',
        'Iluminação: "warm dim light, subtle highlights from candles and pendant lamps"',
    ])

    h2(doc, "Prompts de Vídeo — Sora 2")

    h3(doc, "V-01 — Story natural (15s)")
    code_block(doc,
        "A woman speaking naturally to the camera as if recording\n"
        "a casual Instagram story. She speaks with calm confidence,\n"
        "making small natural hand gestures.\n"
        "The camera is held slightly above eye level, selfie style.\n"
        "Soft indoor lighting. Simple room background.\n"
        "Movement is fluid and organic — not rehearsed.\n"
        "No text on screen. No cuts. Single continuous take."
    )

    h3(doc, "V-02 — Falando enquanto anda")
    code_block(doc,
        "A woman walking casually in an indoor corridor or near a window,\n"
        "glancing at the camera occasionally while speaking softly.\n"
        "Natural steps, relaxed arms, ambient indoor light.\n"
        "Selfie-style framing, slightly unstable as if handheld.\n"
        "Short, candid, real moment feel."
    )

    h3(doc, "V-03 — Expressão / reação estática")
    code_block(doc,
        "A woman reacting with a quiet smile and subtle nod,\n"
        "as if hearing something that resonates with her.\n"
        "No movement beyond the natural expression.\n"
        "Soft light, indoor setting, slight warmth.\n"
        "15 seconds of authentic stillness with micro-expressions."
    )

    h2(doc, "Como usar estes prompts")
    numbered(doc, [
        "Identifique a categoria pelo conteúdo que precisa",
        "Substitua os placeholders pelos dados do Storyboard",
        "Use o LENS para otimizar antes de gerar imagens",
        "Use o REEL para adaptar prompts de vídeo ao Sora 2",
        "Quando aprovar um resultado, salve o prompt nas variações do Bloco 3 do Storyboard",
    ])
    doc.save(f"{OUT_DIR}/08-prompt-pack.docx")
    print("✓ 08-prompt-pack.docx")


# ─────────────────────────────────────────────────────────────────────────────
# 09 — Guia Rápido
# ─────────────────────────────────────────────────────────────────────────────

def doc_09():
    doc = new_doc("09-guia-rapido")
    h1(doc, "Guia Rápido de Referência")
    p(doc, "Bônus — CREATOR$ · O método completo em uma página")

    h2(doc, "O método em 4 passos")

    h3(doc, "Passo 1 — Imagem base")
    bullet(doc, [
        "Ferramenta: Nano Banana Pro ou Higgsfield",
        "Agente: LENS — otimiza o prompt antes de gerar",
        "Meta: 1 imagem-referência aprovada + 20–30 variações curadas",
    ])

    h3(doc, "Passo 2 — Hack da pintura")
    bullet(doc, [
        "Ferramenta: Higgsfield",
        'Prompt: "transforme essa imagem em uma pintura ultra realista com todas as características da personagem, mas traga marcas de tinta visíveis"',
        "Salvar como: pintura-higgsfield.png",
        "Meta: imagem-pintura pronta para o Sora 2",
    ])

    h3(doc, "Passo 3 — Vídeo")
    bullet(doc, [
        "Ferramenta: Sora 2 (via Opera + VPS — instruções no Módulo 3)",
        "Agente: REEL — gera o prompt antes de abrir o Sora",
        "Config: Style = Selfie · Duração = 15s",
        "Input: pintura-higgsfield.png (não a foto original)",
        "Meta: 3 gerações → selecionar o melhor take",
    ])

    h3(doc, "Passo 4 — Curar e organizar")
    bullet(doc, [
        "Pasta: avatar-[nome]/imagens/curadas/  e  videos/curados/",
        "Meta: Storyboard completo preenchido",
    ])

    h2(doc, "Agentes — acesso rápido")
    bullet(doc, [
        "LENS (imagem): https://chatgpt.com/g/g-69be9c381dcc81919e85d56ace38c9f4-lens-r",
        "REEL (vídeo): https://chatgpt.com/g/g-69bebdb35f748191952d6500f350c386-reel-r",
    ])

    h2(doc, "Ferramentas — acesso rápido")
    bullet(doc, [
        "Nano Banana Pro → nanobananapro.com",
        "Higgsfield → higgsfield.ai",
        "Sora 2 → sora.com (via Opera + VPS)",
        "vmake.ai → vmake.ai (remover marca d'água)",
        "CapCut → capcut.com (recortar para Reels 9:16)",
    ])

    h2(doc, "Critérios de seleção")

    h3(doc, "Imagem — aprovar se")
    bullet(doc, [
        "Rosto consistente com a identidade da personagem",
        "Roupa exatamente como descrita no prompt",
        "Iluminação natural (não de estúdio)",
        "Expressão espontânea, não posada",
        "Fundo neutro e sem distrações",
    ])

    h3(doc, "Imagem — descartar se")
    bullet(doc, [
        "Artefatos visíveis (mãos deformadas, objetos duplicados)",
        "Texto ou letras geradas pela IA",
        "Parece rendering 3D ou boneco",
    ])

    h3(doc, "Vídeo — aprovar se")
    bullet(doc, [
        "Movimento fluido (sem teleportes ou artefatos)",
        "Rosto consistente durante todo o vídeo",
        "Roupa não muda no meio do take",
        "Parece gravado, não gerado",
    ])

    h3(doc, "Vídeo — descartar se")
    bullet(doc, [
        "Piscar assimétrico dos olhos",
        "Deformação de pescoço ou mãos",
        "Fundo com elementos que aparecem e somem",
    ])

    h2(doc, "Erros mais comuns")
    bullet(doc, [
        'Parece foto de estúdio → adicione "iPhone photo by a friend, slightly off-center"',
        "Roupa inconsistente → detalhe tecido + modelagem; mova a roupa para mais cedo no prompt",
        "Rosto muda no vídeo → certifique que a pintura-higgsfield.png tem fundo neutro",
        "Movimento rígido no Sora 2 → use sempre Style: Selfie",
        "Aprovando imagens inconsistentes → compare sempre com referencia-base.png",
    ])

    h2(doc, "Checklist de conclusão")

    h3(doc, "Módulo 1")
    bullet(doc, [
        "Worksheets de identidade e voz preenchidos",
        "Storyboard Blocos 1 e 6 completos",
    ])
    h3(doc, "Módulo 2")
    bullet(doc, [
        "Referência-base aprovada e salva",
        "20+ imagens curadas e organizadas",
        "Prompt Mestre no Storyboard (Bloco 3)",
    ])
    h3(doc, "Módulo 3")
    bullet(doc, [
        "Imagem-pintura do Higgsfield salva",
        "3–5 vídeos curados na biblioteca",
        "Prompt de vídeo no Storyboard (Bloco 5)",
    ])
    h3(doc, "Módulo 4")
    bullet(doc, [
        "Storyboard completo (todos os 6 blocos)",
        "Documento final em Notion ou Google Docs",
        "Testado: alguém externo conseguiu descrever o avatar",
    ])
    doc.save(f"{OUT_DIR}/09-guia-rapido.docx")
    print("✓ 09-guia-rapido.docx")


# ─────────────────────────────────────────────────────────────────────────────
# Run all
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    doc_00()
    doc_01()
    doc_02()
    doc_03()
    doc_04()
    doc_05()
    doc_06()
    doc_07()
    doc_08()
    doc_09()
    print("\nDone — all 10 docs generated.")
